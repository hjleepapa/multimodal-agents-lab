#!/usr/bin/env python3
"""
Test DOCX Processing Functionality

This script tests the DOCX processing capabilities of the multimodal agents lab.
"""

import os
from docx import Document
from process_new_data import process_text_file

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Document for Testing', 0)
    
    # Add paragraph
    doc.add_paragraph('This is a sample document created for testing the DOCX processing functionality.')
    
    # Add another paragraph
    doc.add_paragraph('The multimodal agents lab can now process Word documents and extract text from both paragraphs and tables.')
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    
    # Add table data
    row_cells = table.rows[1].cells
    row_cells[0].text = 'PDF Processing'
    row_cells[1].text = 'Extracts images from PDF pages'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'DOCX Processing'
    row_cells[1].text = 'Extracts text from paragraphs and tables'
    
    # Add another paragraph
    doc.add_paragraph('This document demonstrates the ability to process complex Word documents with various formatting elements.')
    
    # Save the document
    os.makedirs('data/text', exist_ok=True)
    doc_path = 'data/text/sample_test.docx'
    doc.save(doc_path)
    
    print(f"✅ Created sample DOCX file: {doc_path}")
    return doc_path

def test_docx_processing():
    """Test DOCX processing functionality"""
    print("🧪 Testing DOCX Processing Functionality")
    print("=" * 50)
    
    # Create sample DOCX file
    docx_path = create_sample_docx()
    
    # Test processing
    print(f"\n📄 Processing DOCX file: {docx_path}")
    try:
        docs = process_text_file(docx_path)
        
        if docs:
            doc = docs[0]
            print(f"✅ Successfully processed DOCX file")
            print(f"📊 Extracted {len(doc['full_content'])} characters")
            print(f"📝 Preview: {doc['content']}")
            print(f"🔍 Full content preview:")
            print("-" * 40)
            print(doc['full_content'][:500] + "..." if len(doc['full_content']) > 500 else doc['full_content'])
            print("-" * 40)
        else:
            print("❌ Failed to process DOCX file")
            
    except Exception as e:
        print(f"❌ Error processing DOCX file: {e}")
        import traceback
        traceback.print_exc()

def test_regular_text_processing():
    """Test regular text file processing for comparison"""
    print(f"\n📝 Testing Regular Text File Processing")
    print("=" * 50)
    
    # Create sample text file
    os.makedirs('data/text', exist_ok=True)
    txt_path = 'data/text/sample_test.txt'
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("This is a sample text file for testing.\n")
        f.write("It contains multiple lines of text.\n")
        f.write("The multimodal agents lab can process both text and DOCX files.\n")
    
    print(f"✅ Created sample text file: {txt_path}")
    
    # Test processing
    print(f"\n📄 Processing text file: {txt_path}")
    try:
        docs = process_text_file(txt_path)
        
        if docs:
            doc = docs[0]
            print(f"✅ Successfully processed text file")
            print(f"📊 Extracted {len(doc['full_content'])} characters")
            print(f"📝 Content: {doc['content']}")
        else:
            print("❌ Failed to process text file")
            
    except Exception as e:
        print(f"❌ Error processing text file: {e}")

def main():
    """Main test function"""
    print("🚀 DOCX Processing Test Suite")
    print("=" * 50)
    
    # Test DOCX processing
    test_docx_processing()
    
    # Test regular text processing
    test_regular_text_processing()
    
    print(f"\n🎉 Test completed!")
    print(f"📁 Check the data/text/ directory for sample files")
    print(f"🔧 You can now add your own DOCX files to data/text/ and run:")
    print(f"   python process_new_data.py")

if __name__ == "__main__":
    main()
